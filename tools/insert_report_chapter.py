"""Script one-off: chèn chương 'VI. Tối ưu hiệu năng realtime' vào báo cáo TTCS.

Thao tác:
1. Chèn câu dẫn vào cuối mục V.5 (Kết quả demo).
2. Chèn toàn bộ chương VI (heading + đoạn văn + 2 bảng) trước chương Định hướng.
3. Đổi heading 'VI. Định hướng phát triển của dự án' thành 'VII. ...'.
4. Thay bullet 'Bổ sung benchmark tốc độ xử lý...' bằng nội dung mới.

Yêu cầu: đã backup file trước khi chạy. Script có guard chống chạy lặp.
"""
from __future__ import annotations

import copy
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = PROJECT_ROOT / "Báo cáo TTCS _ N6 _ CK.docx"

NEW_CHAPTER_TITLE = "VI. Tối ưu hiệu năng realtime"
OLD_CH6_TITLE = "VI. Định hướng phát triển của dự án"
NEW_CH7_TITLE = "VII. Định hướng phát triển của dự án"

V5_LEAD_SENTENCE = (
    "Hiệu năng xử lý chi tiết cùng các kỹ thuật tối ưu realtime của hệ thống "
    "được trình bày ở chương VI."
)

OLD_BENCHMARK_BULLET_PREFIX = "Bổ sung benchmark tốc độ xử lý"
NEW_BENCHMARK_BULLET = (
    "Mở rộng benchmark tốc độ xử lý (FPS, latency end-to-end, mức sử dụng GPU/CPU) "
    "sang các cấu hình phần cứng khác như laptop Intel và embedded board. Kết quả "
    "benchmark trên Apple Silicon đã được trình bày chi tiết ở chương VI; việc bổ "
    "sung số liệu trên các nền tảng còn lại sẽ giúp định lượng đầy đủ tính khả thi "
    "của triển khai thực tế."
)

# ── Nội dung chương VI ───────────────────────────────────────────────────────
# Mỗi phần tử: ("p", style_id, text) hoặc ("table", header_row, data_rows)
CHAPTER_BLOCKS = [
    ("p", "Heading1", NEW_CHAPTER_TITLE),
    ("p", None,
     "Một hệ thống cảnh báo điểm mù chỉ thực sự có giá trị khi cảnh báo được đưa ra "
     "kịp thời. Nếu tốc độ xử lý thấp hơn tốc độ khung hình của camera, thông tin "
     "cảnh báo sẽ đến trễ và mất đi ý nghĩa hỗ trợ người lái. Vì vậy, bên cạnh độ "
     "chính xác của mô hình phát hiện, hiệu năng xử lý theo thời gian thực là yêu "
     "cầu bắt buộc của hệ thống. Chương này trình bày quá trình phân tích điểm "
     "nghẽn hiệu năng, các kỹ thuật tối ưu ở mức ứng dụng, giải pháp tăng tốc suy "
     "luận trên phần cứng Apple Silicon và kết quả benchmark thu được."),

    ("p", "Heading2", "1. Yêu cầu realtime và phân tích bottleneck"),
    ("p", None,
     "Hệ thống đặt ngưỡng realtime tối thiểu là 25 FPS, tương ứng với tốc độ khung "
     "hình phổ biến của camera hành trình và cũng là giá trị target_fps mặc định "
     "trong app.py. Ở cấu hình ban đầu, khi chạy suy luận PyTorch trên CPU của "
     "MacBook M3 (không có GPU CUDA), hệ thống chỉ đạt khoảng 9 FPS, thấp hơn "
     "nhiều so với ngưỡng yêu cầu."),
    ("p", None,
     "Để xác định chính xác điểm nghẽn, nhóm xây dựng công cụ tools/benchmark.py "
     "đo thời gian xử lý của từng bước trong pipeline trên cùng một đoạn video. "
     "Kết quả đo như sau:"),
    ("table",
     ["Bước xử lý", "Thời gian mỗi frame", "Tỷ trọng"],
     [["YOLOv9 inference", "~110 ms", "~95%"],
      ["Tracking + Prediction + Visualization", "< 5 ms", "~5%"]]),
    ("p", None,
     "Kết quả trên cho thấy gần như toàn bộ thời gian xử lý tập trung ở bước suy "
     "luận của mô hình YOLOv9, trong khi các khối tracking, prediction và "
     "visualization chỉ chiếm tỷ trọng không đáng kể. Do đó, chiến lược tối ưu "
     "được chia thành hai hướng bổ trợ nhau: giảm tải ở mức ứng dụng khi hệ thống "
     "không theo kịp thời gian thực, và tăng tốc trực tiếp bước inference bằng "
     "phần cứng chuyên dụng."),

    ("p", "Heading2", "2. Tối ưu mức ứng dụng"),
    ("p", None,
     "Nhóm kỹ thuật thứ nhất hoạt động ở mức ứng dụng, không thay đổi mô hình, "
     "nhằm bảo đảm hệ thống duy trì được nhịp xử lý realtime ngay cả khi phần "
     "cứng không đủ mạnh."),
    ("p", None,
     "Frame skipping với cơ chế Kalman predict-only: khi FPS đo được giảm xuống "
     "dưới 70% ngưỡng mục tiêu, hệ thống chủ động bỏ qua bước detection ở frame "
     "kế tiếp. Điểm quan trọng là các track không bị gián đoạn: nhờ Kalman "
     "Filter, vị trí của từng đối tượng vẫn được ngoại suy bằng bước predict mà "
     "không cần detection mới. Khi hiệu năng phục hồi, detection được bật trở "
     "lại và các track tiếp tục được hiệu chỉnh bằng quan sát thực. Cơ chế này "
     "được kích hoạt qua tham số --enable-frame-skip."),
    ("p", None,
     "Adaptive quality scaling: hệ thống liên tục so sánh FPS đo được với ngưỡng "
     "mục tiêu. Khi FPS thấp hơn 80% target, độ phân giải frame đầu vào được "
     "giảm dần để rút ngắn thời gian suy luận; khi FPS vượt 110% target, độ phân "
     "giải được khôi phục từng bước nhằm giữ chất lượng phát hiện cao nhất có "
     "thể. Cơ chế được kích hoạt qua tham số --enable-adaptive-scale."),
    ("p", None,
     "Bên cạnh đó, ứng dụng hiển thị lớp performance overlay gồm FPS đã làm mượt "
     "và thời gian xử lý của từng giai đoạn, giúp quan sát trực tiếp tác động "
     "của từng kỹ thuật tối ưu trong quá trình chạy. Cần lưu ý rằng hai kỹ thuật "
     "trên là sự đánh đổi có kiểm soát: bỏ qua detection hoặc giảm độ phân giải "
     "đều có thể làm giảm tạm thời chất lượng phát hiện. Vì vậy chúng đóng vai "
     "trò lưới an toàn, còn giải pháp căn cơ vẫn là tăng tốc chính bước "
     "inference."),

    ("p", "Heading2", "3. Tăng tốc inference trên Apple Silicon"),
    ("p", None,
     "Do bottleneck nằm ở bước suy luận, nhóm thực hiện tăng tốc inference trên "
     "nền tảng phần cứng thực tế của nhóm là MacBook M3, theo hai giai đoạn kế "
     "tiếp nhau."),
    ("p", "Heading3", "3.1. Giai đoạn 1: MPS backend"),
    ("p", None,
     "Metal Performance Shaders (MPS) là backend GPU của PyTorch trên Apple "
     "Silicon. Việc chuyển inference từ CPU sang GPU M3 đòi hỏi một số điều "
     "chỉnh trong src/detector.py: hàm select_device của YOLOv9 chỉ nhận biết "
     "CUDA và CPU nên được bypass để khởi tạo trực tiếp torch.device(\"mps\"); "
     "phép NMS chứa toán tử chưa được hỗ trợ trên MPS nên tensor kết quả được "
     "chuyển về CPU trước khi thực hiện NMS; mô hình giữ độ chính xác FP32 vì "
     "MPS đã đủ nhanh và tránh được sai lệch do FP16. Chế độ này được kích hoạt "
     "qua tham số --device mps."),
    ("p", "Heading3", "3.2. Giai đoạn 2: CoreML backend trên Neural Engine"),
    ("p", None,
     "Để khai thác Apple Neural Engine (ANE) — bộ tăng tốc suy luận chuyên dụng "
     "trên chip M3 — mô hình được chuyển đổi sang định dạng CoreML. Quá trình "
     "export được thực hiện bằng công cụ tools/export_coreml.py theo đường "
     "TorchScript → CoreML, sinh ra gói weights/best_roiv2.mlpackage. Ở chế độ "
     "này, YOLOv9Detector bỏ qua hoàn toàn lớp DetectMultiBackend của YOLOv9 và "
     "gọi thẳng CoreML runtime; việc lựa chọn precision do CoreML tự quản lý. "
     "Backend được kích hoạt qua tham số --backend coreml và loại trừ lẫn nhau "
     "với --device mps (hệ thống báo lỗi nếu dùng đồng thời)."),

    ("p", "Heading2", "4. Kết quả benchmark"),
    ("p", None,
     "Bảng dưới đây tổng hợp kết quả FPS đo trên MacBook M3 16GB RAM với video "
     "demo mặc định (assets/videos/demo4.mp4), độ phân giải gốc, không bật frame "
     "skipping và adaptive quality scaling [ĐIỀN ĐIỀU KIỆN ĐO NẾU KHÁC]:"),
    ("table",
     ["Cấu hình", "Backend", "FPS thực đo", "Đạt realtime (≥ 25 FPS)?"],
     [["CPU (baseline)", "PyTorch CPU", "[ĐIỀN SỐ] (~9)", "Không"],
      ["GPU M3", "PyTorch + MPS", "[ĐIỀN SỐ]", "[ĐIỀN]"],
      ["Neural Engine", "CoreML", "55", "Có (≈ 2,2 lần ngưỡng)"]]),
    ("p", None,
     "Với backend CoreML, hệ thống đạt 55 FPS, vượt ngưỡng realtime khoảng 2,2 "
     "lần. Phần dư hiệu năng này có ý nghĩa thực tiễn: nó tạo dư địa cho các mở "
     "rộng trong tương lai như xử lý đồng thời nhiều camera hoặc nâng cấp lên "
     "biến thể YOLOv9 lớn hơn mà vẫn giữ được tính realtime."),

    ("p", "Heading2", "5. Nhận xét và đánh đổi"),
    ("p", None,
     "Kết quả tối ưu cho thấy việc kết hợp hai tầng giải pháp mang lại hiệu quả "
     "rõ rệt: các kỹ thuật mức ứng dụng giúp hệ thống chống chịu khi phần cứng "
     "yếu, trong khi tăng tốc phần cứng giải quyết tận gốc bottleneck inference. "
     "Tuy nhiên, giải pháp hiện tại cũng có giới hạn. CoreML và ANE chỉ khả dụng "
     "trên hệ sinh thái macOS/Apple Silicon, nên kết quả 55 FPS không chuyển "
     "trực tiếp sang các nền tảng khác. Đối với các máy dùng CPU Intel phổ "
     "biến, hướng đi tương đương là OpenVINO kết hợp lượng tử hóa INT8 — phương "
     "án này đã được nhóm thiết kế nhưng chưa triển khai, và được xếp vào định "
     "hướng phát triển ở chương VII. Trong mọi trường hợp, các cơ chế frame "
     "skipping và adaptive quality scaling vẫn được giữ lại như lưới an toàn "
     "cuối cùng, bảo đảm hệ thống không bị tụt nhịp xử lý ngay cả khi chạy trên "
     "phần cứng không có bộ tăng tốc."),
]


def set_paragraph_style(paragraph, style_id: str) -> None:
    """Gán style qua styleId trực tiếp trong XML (tránh phụ thuộc tên style)."""
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:pStyle"))
    if existing is not None:
        p_pr.remove(existing)
    p_style = p_pr.makeelement(qn("w:pStyle"), {qn("w:val"): style_id})
    p_pr.insert(0, p_style)


def replace_paragraph_text(paragraph, new_text: str) -> None:
    """Thay toàn bộ text, giữ run đầu tiên để bảo toàn định dạng."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def find_paragraph(doc, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def build_table(doc, ref_table, header, rows):
    """Tạo bảng mới, clone tblPr từ bảng sẵn có để khớp định dạng báo cáo."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table._tbl.replace(table._tbl.tblPr, copy.deepcopy(ref_table._tbl.tblPr))
    for c, text in enumerate(header):
        cell = table.cell(0, c)
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, text in enumerate(row):
            table.cell(r, c).text = text
    return table


def main() -> int:
    if not DOC_PATH.exists():
        print(f"LỖI: không tìm thấy {DOC_PATH}")
        return 1

    doc = docx.Document(str(DOC_PATH))

    # Guard chống chạy lặp
    if find_paragraph(doc, NEW_CHAPTER_TITLE) is not None:
        print("LỖI: chương VI mới đã tồn tại trong file — không chèn lại.")
        return 1

    anchor = find_paragraph(doc, OLD_CH6_TITLE)
    if anchor is None:
        print(f"LỖI: không tìm thấy heading '{OLD_CH6_TITLE}'.")
        return 1

    n_before = len(doc.paragraphs)

    # 1) Câu dẫn cuối mục V.5 (chèn trước anchor, trước các block chương VI)
    lead = doc.add_paragraph(V5_LEAD_SENTENCE)
    anchor._p.addprevious(lead._p)

    # 2) Chèn các block chương VI theo thứ tự, ngay trước anchor
    ref_table = doc.tables[0]
    n_tables = 0
    for block in CHAPTER_BLOCKS:
        if block[0] == "p":
            _, style_id, text = block
            p = doc.add_paragraph(text)
            if style_id:
                set_paragraph_style(p, style_id)
            anchor._p.addprevious(p._p)
        else:
            _, header, rows = block
            t = build_table(doc, ref_table, header, rows)
            anchor._p.addprevious(t._tbl)
            # đoạn trống sau bảng để Word không dính bảng vào đoạn kế tiếp
            spacer = doc.add_paragraph("")
            anchor._p.addprevious(spacer._p)
            n_tables += 1

    # 3) Đổi số chương Định hướng: VI -> VII
    replace_paragraph_text(anchor, NEW_CH7_TITLE)

    # 4) Thay bullet benchmark trong mục Định hướng §3
    bullet = find_paragraph(doc, OLD_BENCHMARK_BULLET_PREFIX)
    if bullet is None:
        print("LỖI: không tìm thấy bullet benchmark cần thay — hủy, không lưu.")
        return 1
    replace_paragraph_text(bullet, NEW_BENCHMARK_BULLET)

    doc.save(str(DOC_PATH))
    n_para_blocks = sum(1 for b in CHAPTER_BLOCKS if b[0] == "p")
    print(f"OK: đã chèn {n_para_blocks} đoạn + {n_tables} bảng + 1 câu dẫn V.5.")
    print(f"Paragraph body: {n_before} -> {len(doc.paragraphs)}")
    print("Nhớ mở Word và Update Field (F9) để cập nhật mục lục.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
